import json
import time
import uuid
from pathlib import Path
from typing import Dict, Optional, List, Any
import logging
import math
import platform
import re
import subprocess
from datetime import datetime, timezone

import torch

logger = logging.getLogger(__name__)


BENCHMARK_MODE_ORDER = ["baseline", "hybrid_rag", "sgas_no_filtering", "sgas"]

BENCHMARK_MODE_LABELS = {
    "baseline": "Baseline: semantic RAG",
    "hybrid_rag": "Ablation: hybrid RAG",
    "sgas_no_filtering": "Ablation: S-GAS graph ranking only",
    "sgas": "Full S-GAS",
}


class BenchmarkRunner:
    """Benchmark runner for S-GAS system.

    Runs each scenario in four local-evaluation modes:
      1. Baseline: classic semantic RAG only
      2. Hybrid RAG: semantic + BM25/RRF, no graph/swap
      3. S-GAS without filtering: graph/swap path without adaptive filtering
      4. Full S-GAS: graph/swap/filtering enabled

    Compares results and generates a comparative report.
    """

    def __init__(
        self,
        chat_service=None,
        document_processor=None,
        scenarios_dir: str = "tests/scenarios",
        results_dir: str = "logs/benchmarks",
        documents_dir: str = "tests/documents",
    ):
        self._chat_service = chat_service
        self._document_processor = document_processor
        self.scenarios_dir = Path(scenarios_dir)
        self.results_dir = Path(results_dir)
        self.documents_dir = Path(documents_dir)

        self.results_dir.mkdir(parents=True, exist_ok=True)
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        self._benchmark_max_tokens = None
        try:
            from config import Settings
            benchmark_config = Settings().get("benchmark", {})
            self._benchmark_max_tokens = benchmark_config.get("max_tokens")
        except Exception:
            self._benchmark_max_tokens = None

        from .scenario_loader import ScenarioLoader
        from .metrics_collector import MetricsCollector
        from .retrieval_evaluator import RetrievalEvaluator
        from .generation_evaluator import GenerationEvaluator
        from .performance_monitor import PerformanceMonitor

        self.scenario_loader = ScenarioLoader(scenarios_dir)
        self.generation_evaluator = GenerationEvaluator()

        # Separate collectors for each benchmark/ablation mode.
        self._collectors = {mode: MetricsCollector() for mode in BENCHMARK_MODE_ORDER}
        self._retrieval_evaluators = {mode: RetrievalEvaluator() for mode in BENCHMARK_MODE_ORDER}
        self._perf_monitors = {mode: PerformanceMonitor() for mode in BENCHMARK_MODE_ORDER}

        # Backward-compatible aliases for older call sites.
        self._sgas_collector = self._collectors["sgas"]
        self._baseline_collector = self._collectors["baseline"]
        self._sgas_retrieval = self._retrieval_evaluators["sgas"]
        self._baseline_retrieval = self._retrieval_evaluators["baseline"]
        self._sgas_perf = self._perf_monitors["sgas"]
        self._baseline_perf = self._perf_monitors["baseline"]
        self._progress: Dict[str, Dict[str, Any]] = {}

        logger.info("BenchmarkRunner initialized")

    def set_services(self, chat_service, document_processor):
        self._chat_service = chat_service
        self._document_processor = document_processor

    # ── Public API ─────────────────────────────────────────────────────

    def get_progress(self, scenario_name: str) -> Dict[str, Any]:
        return self._progress.get(
            scenario_name,
            {
                "status": "idle",
                "scenario": scenario_name,
                "phase": "idle",
                "percent": 0,
                "message": "No benchmark is running.",
            },
        )

    def _set_progress(
        self,
        scenario_name: str,
        *,
        status: str = "running",
        phase: str,
        percent: float,
        message: str,
        mode: Optional[str] = None,
        turn: int = 0,
        total_turns: int = 0,
    ) -> None:
        self._progress[scenario_name] = {
            "status": status,
            "scenario": scenario_name,
            "phase": phase,
            "mode": mode,
            "turn": turn,
            "total_turns": total_turns,
            "percent": max(0, min(100, round(percent, 1))),
            "message": message,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }

    def _set_mode_progress(
        self,
        scenario_name: str,
        mode_name: str,
        turn: int,
        total_turns: int,
        phase: str = "turn",
    ) -> None:
        mode_index = BENCHMARK_MODE_ORDER.index(mode_name) if mode_name in BENCHMARK_MODE_ORDER else 0
        steps_total = max(1, len(BENCHMARK_MODE_ORDER) * max(1, total_turns))
        steps_done = mode_index * max(1, total_turns) + max(0, turn - 1)
        percent = (steps_done / steps_total) * 100
        self._set_progress(
            scenario_name,
            phase=phase,
            mode=mode_name,
            turn=turn,
            total_turns=total_turns,
            percent=percent,
            message=f"{BENCHMARK_MODE_LABELS.get(mode_name, mode_name)}: turn {turn}/{total_turns}",
        )

    async def _reset_and_flush(self, label: str, fresh_server: bool = False):
        """Reset of stateful services and optionally restart vLLM."""
        if self._chat_service is None:
            return
        self._chat_service.reset_for_benchmark()
        if fresh_server:
            restarted = await self._chat_service._vllm.force_restart()
            if restarted:
                logger.info(f"vLLM server restarted before {label}")
            else:
                logger.warning(f"vLLM restart failed before {label}; continuing with cache flush")
        flushed = await self._chat_service.flush_vllm_cache()
        if flushed:
            logger.info(f"vLLM cache flushed before {label}")

    def _get_mode_spec(self, mode: str) -> Dict[str, Any]:
        """Return benchmark mode behavior without changing app-wide config files."""
        original_bm25 = getattr(self._chat_service, "_bm25_weight", 0.0) if self._chat_service else 0.0
        hybrid_weight = original_bm25 if original_bm25 > 0 else 0.4
        specs = {
            "baseline": {
                "mode": "baseline",
                "label": BENCHMARK_MODE_LABELS["baseline"],
                "baseline_mode": True,
                "bm25_weight": 0.0,
                "enable_sgas_filtering": False,
                "enable_graph_expansion_filter": False,
                "enable_cross_encoder_rerank": False,
                "enable_dynamic_scoring_weights": False,
                "enable_semantic_anchor": False,
                "enable_keyword_boost": False,
                "save_graph": False,
            },
            "hybrid_rag": {
                "mode": "hybrid_rag",
                "label": BENCHMARK_MODE_LABELS["hybrid_rag"],
                "baseline_mode": True,
                "bm25_weight": hybrid_weight,
                "enable_sgas_filtering": False,
                "enable_graph_expansion_filter": False,
                "enable_cross_encoder_rerank": False,
                "enable_dynamic_scoring_weights": False,
                "enable_semantic_anchor": False,
                "enable_keyword_boost": False,
                "save_graph": False,
            },
            "sgas_no_filtering": {
                "mode": "sgas_no_filtering",
                "label": BENCHMARK_MODE_LABELS["sgas_no_filtering"],
                "baseline_mode": False,
                "bm25_weight": original_bm25,
                "enable_sgas_filtering": False,
                "enable_graph_expansion_filter": False,
                "enable_cross_encoder_rerank": False,
                "enable_dynamic_scoring_weights": False,
                "enable_semantic_anchor": False,
                "enable_keyword_boost": False,
                "save_graph": True,
            },
            "sgas": {
                "mode": "sgas",
                "label": BENCHMARK_MODE_LABELS["sgas"],
                "baseline_mode": False,
                "bm25_weight": original_bm25,
                "enable_sgas_filtering": True,
                "enable_graph_expansion_filter": True,
                "enable_cross_encoder_rerank": True,
                "enable_dynamic_scoring_weights": True,
                "enable_semantic_anchor": True,
                "enable_keyword_boost": True,
                "save_graph": True,
            },
        }
        if mode not in specs:
            raise ValueError(f"Invalid benchmark mode: {mode}")
        return specs[mode]

    def _apply_mode_overrides(self, mode_spec: Dict[str, Any]) -> Dict[str, Any]:
        """Temporarily tune ChatService for one benchmark/ablation mode."""
        if self._chat_service is None:
            return {}
        saved = {
            "bm25_weight": getattr(self._chat_service, "_bm25_weight", 0.0),
            "enable_sgas_filtering": getattr(self._chat_service, "_enable_sgas_filtering", True),
            "enable_graph_expansion_filter": getattr(self._chat_service, "_enable_graph_expansion_filter", True),
            "enable_cross_encoder_rerank": getattr(self._chat_service, "_enable_cross_encoder_rerank", True),
            "enable_dynamic_scoring_weights": getattr(self._chat_service, "_enable_dynamic_scoring_weights", True),
            "enable_semantic_anchor": getattr(self._chat_service, "_enable_semantic_anchor", True),
            "enable_keyword_boost": getattr(self._chat_service, "_enable_keyword_boost", True),
            "exclude_used_chunks": getattr(self._chat_service, "_exclude_used_chunks", False),
        }
        self._chat_service._bm25_weight = mode_spec["bm25_weight"]
        self._chat_service._enable_sgas_filtering = mode_spec["enable_sgas_filtering"]
        self._chat_service._enable_graph_expansion_filter = mode_spec["enable_graph_expansion_filter"]
        self._chat_service._enable_cross_encoder_rerank = mode_spec["enable_cross_encoder_rerank"]
        self._chat_service._enable_dynamic_scoring_weights = mode_spec["enable_dynamic_scoring_weights"]
        self._chat_service._enable_semantic_anchor = mode_spec["enable_semantic_anchor"]
        self._chat_service._enable_keyword_boost = mode_spec["enable_keyword_boost"]
        self._chat_service._exclude_used_chunks = False
        return saved

    def _restore_mode_overrides(self, saved: Dict[str, Any]) -> None:
        if self._chat_service is None or not saved:
            return
        self._chat_service._bm25_weight = saved["bm25_weight"]
        self._chat_service._enable_sgas_filtering = saved["enable_sgas_filtering"]
        self._chat_service._enable_graph_expansion_filter = saved["enable_graph_expansion_filter"]
        self._chat_service._enable_cross_encoder_rerank = saved["enable_cross_encoder_rerank"]
        self._chat_service._enable_dynamic_scoring_weights = saved["enable_dynamic_scoring_weights"]
        self._chat_service._enable_semantic_anchor = saved["enable_semantic_anchor"]
        self._chat_service._enable_keyword_boost = saved["enable_keyword_boost"]
        self._chat_service._exclude_used_chunks = saved["exclude_used_chunks"]

    async def run_scenario(
        self,
        scenario_name: str,
        sessions: dict,
        document_processor=None,
        session_id: Optional[str] = None,
        fresh_server: bool = False,
    ) -> Dict:
        """Run scenario in baseline, hybrid RAG, S-GAS ablation, and full S-GAS modes."""
        doc_processor = document_processor or self._document_processor
        if doc_processor is None:
            raise ValueError("document_processor is required")

        scenario = self.scenario_loader.load_scenario(scenario_name)
        self._set_progress(
            scenario_name,
            phase="prepare",
            percent=1,
            message=f"Preparing {scenario.name}: loading scenario and document",
            total_turns=len(scenario.turns),
        )
        logger.info(f"=== Benchmark: {scenario.name} ({len(scenario.turns)} turns) ===")
        logger.info(f"Document: {scenario.document}")

        doc_path = self.documents_dir / scenario.document
        if not doc_path.exists():
            raise FileNotFoundError(f"Document not found: {doc_path}")

        mode_results: Dict[str, Dict] = {}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        for mode in BENCHMARK_MODE_ORDER:
            mode_spec = self._get_mode_spec(mode)
            self._set_mode_progress(scenario_name, mode, 1, len(scenario.turns), phase="mode_start")
            await self._reset_and_flush(f"{mode_spec['label']} run", fresh_server=fresh_server)
            logger.info(f"── Mode: {mode_spec['label']} ──")

            saved_overrides = self._apply_mode_overrides(mode_spec)
            mode_session = self._create_session(sessions)
            try:
                self._upload_document(mode_session, doc_path, doc_processor,
                                      chunk_size=scenario.chunk_size,
                                      chunk_overlap=scenario.chunk_overlap)
                mode_result = await self._run_mode(
                    scenario=scenario, sessions=sessions, session_id=mode_session,
                    mode_spec=mode_spec, collector=self._collectors[mode],
                    retrieval_eval=self._retrieval_evaluators[mode],
                    perf_monitor=self._perf_monitors[mode],
                    run_id=uuid.uuid4().hex,
                    progress_scenario_name=scenario_name,
                )
            finally:
                self._restore_mode_overrides(saved_overrides)
                self._cleanup_session(mode_session)

            if mode_spec["save_graph"]:
                self._save_graph_snapshot(
                    scenario_name,
                    timestamp=f"{mode}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                )

            csv_path = self.results_dir / f"{scenario_name}_{mode}_{timestamp}.csv"
            json_path = self.results_dir / f"{scenario_name}_{mode}_{timestamp}.json"
            self._collectors[mode].export_to_csv(str(csv_path))
            self._collectors[mode].export_to_json(str(json_path))
            mode_results[mode] = {
                **mode_result,
                "label": mode_spec["label"],
                "csv_file": str(csv_path),
                "json_file": str(json_path),
            }

        self._set_progress(
            scenario_name,
            phase="report",
            percent=96,
            message="Generating comparison JSON and DOCX summary report",
            total_turns=len(scenario.turns),
        )
        comparison = self._compare_mode_results({m: r["summary"] for m, r in mode_results.items()})
        comparison_json = self.results_dir / f"{scenario_name}_comparison_{timestamp}.json"
        docx_report = self.results_dir / f"{scenario_name}_summary_report_{timestamp}.docx"

        with open(comparison_json, 'w', encoding='utf-8') as f:
            json.dump(comparison, f, indent=2, ensure_ascii=False, default=str)

        self._generate_docx_summary_report(
            scenario=scenario,
            mode_results=mode_results,
            comparison=comparison,
            output_path=docx_report,
        )

        logger.info(f"Benchmark completed! Comparison: {comparison_json}")
        logger.info(f"Benchmark DOCX report: {docx_report}")
        self._log_comparison(comparison)
        overall_status = (
            "success"
            if all(r.get("summary", {}).get("status") == "completed" for r in mode_results.values())
            else "partial_failed"
        )

        self._set_progress(
            scenario_name,
            status=overall_status,
            phase="completed",
            percent=100,
            message="Benchmark completed" if overall_status == "success" else "Benchmark completed with partial failures",
            total_turns=len(scenario.turns),
        )

        return {
            'status': overall_status,
            'scenario': scenario.name,
            'modes': mode_results,
            # Backward-compatible aliases.
            'sgas': mode_results['sgas'],
            'baseline': mode_results['baseline'],
            'comparison': comparison,
            'comparison_file': str(comparison_json),
            'docx_report_file': str(docx_report),
        }

    async def run_single_mode(
        self,
        scenario_name: str,
        sessions: dict,
        mode: str = "sgas",
        document_processor=None,
        fresh_server: bool = False,
    ) -> Dict:
        """Running scenario in a single benchmark/ablation mode only."""
        doc_processor = document_processor or self._document_processor
        if doc_processor is None:
            raise ValueError("document_processor is required")

        mode_spec = self._get_mode_spec(mode)
        scenario = self.scenario_loader.load_scenario(scenario_name)
        self._set_progress(
            scenario_name,
            phase="prepare",
            mode=mode,
            percent=1,
            message=f"Preparing {scenario.name}: {mode_spec['label']}",
            total_turns=len(scenario.turns),
        )
        logger.info(f"=== Single-mode Benchmark: {scenario.name} ({mode}) ===")

        doc_path = self.documents_dir / scenario.document
        if not doc_path.exists():
            raise FileNotFoundError(f"Document not found: {doc_path}")

        await self._reset_and_flush(f"{mode} run", fresh_server=fresh_server)

        run_id = uuid.uuid4().hex
        collector = self._collectors[mode]
        retrieval_eval = self._retrieval_evaluators[mode]
        perf_monitor = self._perf_monitors[mode]

        saved_overrides = self._apply_mode_overrides(mode_spec)
        session_id = self._create_session(sessions)
        try:
            self._upload_document(session_id, doc_path, doc_processor,
                                  chunk_size=scenario.chunk_size,
                                  chunk_overlap=scenario.chunk_overlap)
            result = await self._run_mode(
                scenario=scenario, sessions=sessions, session_id=session_id,
                mode_spec=mode_spec, collector=collector,
                retrieval_eval=retrieval_eval, perf_monitor=perf_monitor,
                run_id=run_id,
                progress_scenario_name=scenario_name,
            )
        finally:
            self._restore_mode_overrides(saved_overrides)
            self._cleanup_session(session_id)

        if mode_spec["save_graph"]:
            self._save_graph_snapshot(scenario_name, timestamp=datetime.now().strftime("%Y%m%d_%H%M%S"))

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        csv_path = self.results_dir / f"{scenario_name}_{mode}_{timestamp}.csv"
        json_path = self.results_dir / f"{scenario_name}_{mode}_{timestamp}.json"
        collector.export_to_csv(str(csv_path))
        collector.export_to_json(str(json_path))
        self._set_progress(
            scenario_name,
            status=result['summary'].get('status', 'completed'),
            phase="completed",
            mode=mode,
            percent=100,
            message=f"{mode_spec['label']} completed",
            total_turns=len(scenario.turns),
        )

        return {
            'status': result['summary'].get('status', 'completed'),
            'scenario': scenario.name,
            'mode': mode,
            'label': mode_spec['label'],
            'session_id': session_id,
            'summary': result['summary'],
            'csv_file': str(csv_path),
            'json_file': str(json_path),
        }

    # ── Run single mode ────────────────────────────────────────────────

    async def _run_mode(
        self, scenario, sessions, session_id, mode_spec,
        collector, retrieval_eval, perf_monitor,
        run_id: Optional[str] = None,
        progress_scenario_name: Optional[str] = None,
    ) -> Dict:
        mode_name = mode_spec["mode"]
        baseline_mode = mode_spec["baseline_mode"]

        collector.start_session()
        retrieval_eval.reset()
        perf_monitor.reset()

        # Measuring VRAM baseline before any work
        vram_baseline_gb = 0.0
        if torch.cuda.is_available():
            torch.cuda.synchronize()
            vram_baseline_gb = torch.cuda.memory_allocated() / (1024 ** 3)
            logger.info(f"  [{mode_name}] VRAM baseline: {vram_baseline_gb:.4f} GB (subtracted from measurements)")

        questions, answers = [], []
        retrieved_chunks_list, ground_truth_chunks_list, ground_truth_answers = [], [], []
        retrieved_texts_list, ground_truth_texts_list = [], []
        evidence_texts_list = []
        total_chunks_per_turn = []
        failed_turn = None
        failure_error = ""

        for i, turn in enumerate(scenario.turns):
            if progress_scenario_name:
                self._set_mode_progress(
                    progress_scenario_name,
                    mode_name,
                    i + 1,
                    len(scenario.turns),
                )
            logger.info(f"  [{mode_name}] Turn {i + 1}/{len(scenario.turns)}: {turn.query[:60]}...")

            gt_chunks = turn.ground_truth if turn.ground_truth else []
            gt_answer = turn.ground_truth_answer if turn.ground_truth_answer else ""
            gt_text = turn.ground_truth_text if turn.ground_truth_text else ""
            evidence_texts = self._get_turn_evidence_texts(turn)

            perf_monitor.start_timing('turn')
            try:
                result = await self._run_turn(
                    session_id=session_id, sessions=sessions, query=turn.query,
                    ground_truth_chunks=gt_chunks, ground_truth_answer=gt_answer,
                    baseline_mode=baseline_mode, run_id=run_id,
                )
            except Exception as exc:
                failed_turn = i + 1
                failure_error = str(exc)
                logger.error(
                    "  [%s] Turn %s failed; saving partial benchmark data and stopping this mode: %s",
                    mode_name,
                    failed_turn,
                    failure_error,
                )
                result = self._failed_turn_result(
                    query=turn.query,
                    ground_truth_chunks=gt_chunks,
                    ground_truth_answer=gt_answer,
                    error=failure_error,
                    elapsed_ms=perf_monitor.end_timing('turn') * 1000,
                )
                collector.record_turn(result)
                questions.append(turn.query)
                answers.append("")
                retrieved_chunks_list.append([])
                ground_truth_chunks_list.append(gt_chunks)
                ground_truth_answers.append(gt_answer)
                retrieved_texts_list.append([])
                ground_truth_texts_list.append(gt_text)
                evidence_texts_list.append(evidence_texts)
                total_chunks_per_turn.append(20)
                await self._recover_vllm_after_failure(mode_name, failed_turn)
                break

            # Subtracting VRAM baseline
            result['vram_allocated_gb'] = max(0.0, result['vram_allocated_gb'] - vram_baseline_gb)
            result['vram_reserved_gb'] = max(0.0, result['vram_reserved_gb'] - vram_baseline_gb)
            result['vram_peak_gb'] = max(0.0, result['vram_peak_gb'] - vram_baseline_gb)

            norm_retrieved = self._normalize_chunk_ids(result.get('retrieved_chunks', []))
            norm_gt = self._normalize_chunk_ids(gt_chunks)
            turn_retrieval = retrieval_eval.evaluate_retrieval(
                retrieved_chunks=norm_retrieved,
                relevant_chunks=norm_gt,
                used_chunks=set(),
                total_chunks=result.get('total_chunks', 20) or 20,
                top_k_values=[5, 10, 20],
                record=False,
            )
            result.update(self._flatten_turn_retrieval_metrics(turn_retrieval))
            evidence_metrics = self._evaluate_evidence_for_turn(
                retrieved_texts=result.get('retrieved_chunk_texts', []),
                evidence_texts=evidence_texts,
                k=5,
            )
            result.update(evidence_metrics)

            if gt_answer:
                turn_generation = self.generation_evaluator.evaluate_generation(
                    generated_answer=result.get('response', ''),
                    reference_answer=gt_answer,
                )
                rouge = turn_generation.get('rouge_scores', {})
                result.update({
                    'answer_semantic_similarity': turn_generation.get('bertscore', 0),
                    'answer_token_f1': turn_generation.get('token_f1', 0),
                    'answer_exact_match': turn_generation.get('exact_match', 0),
                    'answer_rouge1': rouge.get('rouge1', 0),
                    'answer_rouge2': rouge.get('rouge2', 0),
                    'answer_rougeL': rouge.get('rougeL', 0),
                })

            collector.record_turn(result)
            questions.append(turn.query)
            answers.append(result['response'])
            retrieved_chunks_list.append(result.get('retrieved_chunks', []))
            ground_truth_chunks_list.append(gt_chunks)
            ground_truth_answers.append(gt_answer)
            retrieved_texts_list.append(result.get('retrieved_chunk_texts', []))
            ground_truth_texts_list.append(gt_text)
            evidence_texts_list.append(evidence_texts)
            total_chunks_per_turn.append(result.get('total_chunks', 20))

            logger.info(f"    Coverage: {result['coverage_ratio'] * 100:.1f}%")
            logger.info(f"    Latency: {result['latency_ms']:.1f} ms "
                        f"(search={result.get('latency_search_ms', 0):.0f}, "
                        f"rerank={result.get('latency_rerank_ms', 0):.0f}, "
                        f"swap={result.get('latency_swap_ms', 0):.0f}, "
                        f"inference={result.get('latency_inference_ms', 0):.0f})")
            logger.info(f"    VRAM: alloc={result['vram_allocated_gb']:.3f} GB, peak={result['vram_peak_gb']:.3f} GB")
            logger.info(f"    RAM: {result.get('ram_used_gb', 0):.2f} GB ({result.get('ram_percent', 0):.1f}%), "
                        f"Disk: {result.get('disk_used_gb', 0):.1f} GB ({result.get('disk_percent', 0):.1f}%), "
                        f"Process RSS: {result.get('process_rss_mb', 0):.1f} MB")
            logger.info(f"    Graph: {result['graph_nodes']} nodes, {result['graph_edges']} edges")
            logger.info(f"    Swap: {result['swap_operations']} ops, Cache hit: {result['cache_hit_rate']:.3f}")

            perf_monitor.end_timing('turn')
            perf_monitor.record_vram(result['vram_allocated_gb'], result['vram_reserved_gb'])

        summary = collector.calculate_summary()
        summary['mode'] = mode_name
        summary['mode_label'] = mode_spec['label']
        if failed_turn is not None:
            summary['status'] = 'failed' if summary.get('total_turns', 0) == 0 else 'partial_failed'
            summary['failed_turn'] = failed_turn
            summary['error'] = failure_error
        else:
            summary['status'] = 'completed'

        # Retrieval evaluation
        retrieval_results = self._evaluate_retrieval(
            retrieved_chunks_list, ground_truth_chunks_list, retrieval_eval,
            total_chunks_per_turn=total_chunks_per_turn,
            retrieved_texts_list=retrieved_texts_list,
            ground_truth_texts_list=ground_truth_texts_list,
            evidence_texts_list=evidence_texts_list,
        )
        summary['retrieval_metrics'] = retrieval_results

        # Generation evaluation
        if any(ground_truth_answers):
            generation_results = self.generation_evaluator.evaluate_multi_turn(
                generated_answers=answers, reference_answers=ground_truth_answers,
            )
            summary['generation_metrics'] = generation_results

        return {'session_id': session_id, 'summary': summary}

    async def _recover_vllm_after_failure(self, mode_name: str, failed_turn: int) -> None:
        if self._chat_service is None:
            return
        try:
            health = await self._chat_service._vllm.check_health()
            if health != "healthy":
                logger.warning(
                    "vLLM is %s after %s turn %s; attempting managed restart before the next task",
                    health,
                    mode_name,
                    failed_turn,
                )
                await self._chat_service._vllm.force_restart()
        except Exception as exc:
            logger.warning("vLLM recovery check failed: %s", exc)

    @staticmethod
    def _failed_turn_result(
        query: str,
        ground_truth_chunks: List[str],
        ground_truth_answer: str,
        error: str,
        elapsed_ms: float = 0,
    ) -> Dict:
        return {
            'status': 'failed',
            'error': error,
            'query': query,
            'response': '',
            'latency_ms': elapsed_ms,
            'latency_search_ms': 0,
            'latency_rerank_ms': 0,
            'latency_swap_ms': 0,
            'latency_inference_ms': elapsed_ms,
            'vram_allocated_gb': 0,
            'vram_reserved_gb': 0,
            'vram_peak_gb': 0,
            'ram_used_gb': 0,
            'ram_percent': 0,
            'disk_used_gb': 0,
            'disk_percent': 0,
            'process_rss_mb': 0,
            'active_chunks': 0,
            'total_chunks': 0,
            'coverage_ratio': 0,
            'cache_hit_rate': 0,
            'cache_hits': 0,
            'cache_misses': 0,
            'swap_operations': 0,
            'graph_nodes': 0,
            'graph_edges': 0,
            'retrieved_chunks': [],
            'retrieved_chunk_texts': [],
            'chunk_scores': [],
            'ground_truth_chunks': ground_truth_chunks,
            'ground_truth_answer': ground_truth_answer,
        }

    # ── Turn execution ─────────────────────────────────────────────────

    async def _run_turn(
        self, session_id: str, sessions: dict, query: str,
        ground_truth_chunks: List[str], ground_truth_answer: str,
        baseline_mode: bool = False, run_id: Optional[str] = None,
    ) -> Dict:
        """Run single turn through real S-GAS pipeline via ChatService."""
        session_data = sessions[session_id]

        if torch.cuda.is_available():
            torch.cuda.reset_peak_memory_stats()

        start_time = time.time()

        if self._chat_service is not None:
            chat_result = await self._chat_service.process_chat(
                session_id=session_id, session_data=session_data,
                message=query, use_rag=True, n_chunks=5, temperature=0.1,
                max_tokens=self._benchmark_max_tokens,
                baseline_mode=baseline_mode, run_id=run_id,
            )

            metadata = chat_result.get('metadata', {})
            swap_stats = metadata.get('swap_statistics', {})
            latency_info = metadata.get('latency', {})
            vram_info = metadata.get('vram', {})
            graph_stats = metadata.get('graph_statistics', {})
            cache_check = metadata.get('cache_check', {})
            vllm_obs = metadata.get('vllm_observability', {})
            vllm_after = vllm_obs.get('after', {})
            gpu_window = vllm_obs.get('gpu_inference_window', {})

            vram_after = vram_info.get('after_inference', {})

            sys_res = metadata.get('system_resources', {})

            return {
                'query': query,
                'response': chat_result.get('response', ''),
                # Latency breakdown
                'latency_ms': latency_info.get('total_ms', (time.time() - start_time) * 1000),
                'latency_search_ms': latency_info.get('search_ms', 0),
                'latency_rerank_ms': latency_info.get('rerank_ms', 0),
                'latency_swap_ms': latency_info.get('swap_ms', 0),
                'latency_inference_ms': latency_info.get('inference_ms', 0),
                'latency_observability_ms': latency_info.get('observability_ms', 0),
                # VRAM
                'vram_allocated_gb': vram_after.get('allocated_gb', 0),
                'vram_reserved_gb': vram_after.get('reserved_gb', 0),
                'vram_peak_gb': vram_after.get('peak_gb', 0),
                # System resources
                'ram_used_gb': sys_res.get('ram', {}).get('used_gb', 0),
                'ram_percent': sys_res.get('ram', {}).get('percent', 0),
                'disk_used_gb': sys_res.get('disk', {}).get('used_gb', 0),
                'disk_percent': sys_res.get('disk', {}).get('percent', 0),
                'process_rss_mb': sys_res.get('process', {}).get('rss_mb', 0),
                # vLLM /metrics and physical GPU inference-window samples
                'vllm_metrics_available': vllm_obs.get('available', False),
                'vllm_kv_cache_usage_before': vllm_obs.get('kv_cache_usage_before', 0) or 0,
                'vllm_kv_cache_usage_after': vllm_obs.get('kv_cache_usage_after', 0) or 0,
                'vllm_kv_cache_usage_delta': vllm_obs.get('kv_cache_usage_delta', 0) or 0,
                'vllm_prefix_cache_hit_rate_delta': vllm_obs.get('prefix_cache_hit_rate_delta', 0),
                'vllm_prefix_cache_queries_delta': vllm_obs.get('prefix_cache_queries_delta', 0),
                'vllm_prefix_cache_hits_delta': vllm_obs.get('prefix_cache_hits_delta', 0),
                'vllm_preemptions_delta': vllm_obs.get('num_preemptions_delta', 0),
                'vllm_request_success_delta': vllm_obs.get('request_success_delta', 0),
                'vllm_generation_tokens_delta': vllm_obs.get('generation_tokens_delta', 0),
                'vllm_prompt_tokens_delta': vllm_obs.get('prompt_tokens_delta', 0),
                'vllm_tokens_per_second': vllm_obs.get('tokens_per_second', 0),
                'vllm_total_tokens_per_second': vllm_obs.get('total_tokens_per_second', 0),
                'vllm_ttft_avg_s': vllm_after.get('ttft_avg_s') or 0,
                'vllm_inter_token_latency_avg_s': vllm_after.get('inter_token_latency_avg_s') or 0,
                'vllm_prefill_time_avg_s': vllm_after.get('prefill_time_avg_s') or 0,
                'vllm_decode_time_avg_s': vllm_after.get('decode_time_avg_s') or 0,
                'vllm_queue_time_avg_s': vllm_after.get('queue_time_avg_s') or 0,
                'gpu_sample_count': gpu_window.get('gpu_sample_count', 0),
                'gpu_utilization_avg_pct': gpu_window.get('gpu_utilization_avg_pct', 0),
                'gpu_utilization_peak_pct': gpu_window.get('gpu_utilization_peak_pct', 0),
                'gpu_memory_used_peak_mb': gpu_window.get('gpu_memory_used_peak_mb', 0),
                'gpu_memory_used_peak_pct': gpu_window.get('gpu_memory_used_peak_pct', 0),
                'gpu_power_avg_w': gpu_window.get('gpu_power_avg_w', 0),
                'gpu_power_peak_w': gpu_window.get('gpu_power_peak_w', 0),
                # Chunks
                'active_chunks': metadata.get('context_chunks_used', 0),
                'total_chunks': metadata.get('total_chunks_available', 0),
                'coverage_ratio': metadata.get('coverage_percent', 0) / 100.0,
                # Cache / Swap
                'cache_hit_rate': cache_check.get('hit_rate', 0),
                'cache_hits': cache_check.get('hits', 0),
                'cache_misses': cache_check.get('misses', 0),
                'swap_operations': swap_stats.get('total_swap_operations', 0),
                # Graph
                'graph_nodes': graph_stats.get('total_nodes', 0),
                'graph_edges': graph_stats.get('total_edges', 0),
                # Retrieval
                'retrieved_chunks': metadata.get('retrieved_chunk_ids', []),
                'retrieved_chunk_texts': metadata.get('retrieved_chunk_texts', []),
                'chunk_scores': metadata.get('chunk_scores', []),
                # Ground truth
                'ground_truth_chunks': ground_truth_chunks,
                'ground_truth_answer': ground_truth_answer,
            }
        else:
            raise RuntimeError(
                "BenchmarkRunner requires a real chat_service. "
                "Mock responses are forbidden in benchmarks because they would "
                "produce fake metrics. Initialize BenchmarkRunner with a valid "
                "ChatService instance before running scenarios."
            )

    # ── Evaluation helpers ─────────────────────────────────────────────

    @staticmethod
    def _normalize_chunk_ids(chunk_ids: List[str]) -> List[str]:
        """Normalizing chunk IDs to comparable chunk-index form."""
        normalized = []
        for cid in chunk_ids:
            # Runtime format: "session_abc:doc-uuid:3"
            if ':' in cid:
                normalized.append(cid.rsplit(':', 1)[-1])
            # Scenario format: "chunk_001" → "0", "chunk_023" → "22"
            elif cid.startswith('chunk_'):
                try:
                    normalized.append(str(int(cid.split('_', 1)[1]) - 1))
                except ValueError:
                    normalized.append(cid)
            else:
                normalized.append(cid)
        return normalized

    @staticmethod
    def _flatten_turn_retrieval_metrics(metrics: Dict) -> Dict:
        """Pick stable per-turn retrieval metrics for CSV export."""
        return {
            'retrieval_recall_at_5': metrics.get('recall_at_k', {}).get('recall@5', 0),
            'retrieval_precision_at_5': metrics.get('precision_at_k', {}).get('precision@5', 0),
            'retrieval_f1_at_5': metrics.get('f1_at_k', {}).get('f1@5', 0),
            'retrieval_hit_at_5': metrics.get('hit_at_k', {}).get('hit@5', 0),
            'retrieval_mrr': metrics.get('mrr', 0),
            'retrieval_ndcg_at_5': metrics.get('ndcg_at_k', {}).get('ndcg@5', 0),
            'retrieval_map_at_5': metrics.get('average_precision_at_k', {}).get('ap@5', 0),
            'retrieval_relevant_count': metrics.get('relevant_count', 0),
            'retrieval_retrieved_count': metrics.get('retrieved_count', 0),
            # Legacy columns kept meaningful for existing reports.
            'recall_at_k': metrics.get('recall_at_k', {}).get('recall@5', 0),
            'precision': metrics.get('precision_at_k', {}).get('precision@5', 0),
        }

    @staticmethod
    def _normalize_text(text: str) -> str:
        return " ".join(re.findall(r"\w+", (text or "").lower(), flags=re.UNICODE))

    @classmethod
    def _token_f1(cls, left: str, right: str) -> float:
        left_tokens = cls._normalize_text(left).split()
        right_tokens = cls._normalize_text(right).split()
        if not left_tokens or not right_tokens:
            return 0.0
        right_counts = {}
        for token in right_tokens:
            right_counts[token] = right_counts.get(token, 0) + 1
        overlap = 0
        for token in left_tokens:
            if right_counts.get(token, 0) > 0:
                overlap += 1
                right_counts[token] -= 1
        if overlap == 0:
            return 0.0
        precision = overlap / len(left_tokens)
        recall = overlap / len(right_tokens)
        return 2 * precision * recall / (precision + recall)

    @staticmethod
    def _get_turn_evidence_texts(turn) -> List[str]:
        """Return evidence snippets, supporting both new list and legacy single phrase."""
        evidence = []
        if getattr(turn, 'ground_truth_texts', None):
            evidence.extend([t for t in turn.ground_truth_texts if t])
        if getattr(turn, 'ground_truth_text', ''):
            evidence.append(turn.ground_truth_text)
        deduped = []
        seen = set()
        for text in evidence:
            key = text.strip().lower()
            if key and key not in seen:
                deduped.append(text.strip())
                seen.add(key)
        return deduped

    @classmethod
    def _evidence_matches_chunk(cls, evidence: str, chunk_text: str, fuzzy_threshold: float = 0.65) -> bool:
        """Evidence match is exact normalized containment or strong token-F1."""
        ev_norm = cls._normalize_text(evidence)
        chunk_norm = cls._normalize_text(chunk_text)
        if not ev_norm or not chunk_norm:
            return False
        if ev_norm in chunk_norm:
            return True
        return cls._token_f1(evidence, chunk_text) >= fuzzy_threshold

    @classmethod
    def _evaluate_evidence_for_turn(
        cls,
        retrieved_texts: List[str],
        evidence_texts: List[str],
        k: int = 5,
    ) -> Dict:
        """Chunk-ID-independent evidence metrics for one turn."""
        if not evidence_texts:
            return {
                'evidence_recall_at_5': 0.0,
                'evidence_hit_at_5': 0.0,
                'evidence_mrr': 0.0,
                'evidence_ndcg_at_5': 0.0,
                'evidence_map_at_5': 0.0,
                'evidence_token_f1_at_5': 0.0,
                'evidence_count': 0,
            }

        top_texts = retrieved_texts[:k]
        matched_evidence = set()
        first_match_rank = None
        chunk_relevance = []
        precision_sum = 0.0
        hits = 0
        best_f1_scores = []

        for evidence_idx, evidence in enumerate(evidence_texts):
            best_f1 = 0.0
            for chunk_text in top_texts:
                best_f1 = max(best_f1, cls._token_f1(evidence, chunk_text))
            best_f1_scores.append(best_f1)

            for rank, chunk_text in enumerate(top_texts, start=1):
                if cls._evidence_matches_chunk(evidence, chunk_text):
                    matched_evidence.add(evidence_idx)
                    if first_match_rank is None or rank < first_match_rank:
                        first_match_rank = rank
                    break

        seen_any = set()
        for rank, chunk_text in enumerate(top_texts, start=1):
            relevant_here = False
            for evidence_idx, evidence in enumerate(evidence_texts):
                if evidence_idx in seen_any:
                    continue
                if cls._evidence_matches_chunk(evidence, chunk_text):
                    relevant_here = True
                    seen_any.add(evidence_idx)
            chunk_relevance.append(1.0 if relevant_here else 0.0)
            if relevant_here:
                hits += 1
                precision_sum += hits / rank

        evidence_count = len(evidence_texts)
        recall = len(matched_evidence) / evidence_count
        hit = 1.0 if matched_evidence else 0.0
        mrr = 1.0 / first_match_rank if first_match_rank else 0.0
        dcg = sum(rel / math.log2(idx + 2) for idx, rel in enumerate(chunk_relevance))
        ideal_hits = min(evidence_count, k)
        idcg = sum(1.0 / math.log2(idx + 2) for idx in range(ideal_hits))
        ndcg = dcg / idcg if idcg > 0 else 0.0
        ap = precision_sum / ideal_hits if ideal_hits > 0 else 0.0
        avg_token_f1 = sum(best_f1_scores) / len(best_f1_scores) if best_f1_scores else 0.0

        return {
            'evidence_recall_at_5': recall,
            'evidence_hit_at_5': hit,
            'evidence_mrr': mrr,
            'evidence_ndcg_at_5': ndcg,
            'evidence_map_at_5': ap,
            'evidence_token_f1_at_5': avg_token_f1,
            'evidence_count': evidence_count,
        }

    def _evaluate_retrieval(
        self, retrieved_list: List[List[str]], gt_list: List[List[str]],
        evaluator, total_chunks_per_turn: List[int] = None,
        retrieved_texts_list: List[List[str]] = None,
        ground_truth_texts_list: List[str] = None,
        evidence_texts_list: List[List[str]] = None,
    ) -> Dict:
        all_results = []
        for i, (retrieved, ground_truth) in enumerate(zip(retrieved_list, gt_list)):
            norm_retrieved = self._normalize_chunk_ids(retrieved)
            norm_gt = self._normalize_chunk_ids(ground_truth)
            total = (total_chunks_per_turn[i] if total_chunks_per_turn else None) or 20
            result = evaluator.evaluate_retrieval(
                retrieved_chunks=norm_retrieved, relevant_chunks=norm_gt,
                used_chunks=set(), total_chunks=total, top_k_values=[5, 10, 20],
            )
            all_results.append(result)

        valid_results = [r for r in all_results if r['relevant_count'] > 0]
        n = len(valid_results) if valid_results else 1

        def _avg_nested(section: str, key: str) -> float:
            return sum(r.get(section, {}).get(key, 0.0) for r in valid_results) / n

        def _avg_flat(key: str) -> float:
            return sum(r.get(key, 0.0) for r in valid_results) / n

        # ── Text Recall: chunk-ID-independent ────────────────────────
        # Binary per-turn: 1.0 if ANY retrieved chunk contains the ground_truth_text phrase, else 0.0.
        text_recall_scores = []
        if retrieved_texts_list and ground_truth_texts_list:
            for turn_texts, gt_phrase in zip(retrieved_texts_list, ground_truth_texts_list):
                if not gt_phrase:
                    continue
                phrase_lower = gt_phrase.lower()
                combined = ' '.join(turn_texts).lower()
                text_recall_scores.append(1.0 if phrase_lower in combined else 0.0)

        avg_text_recall = (
            sum(text_recall_scores) / len(text_recall_scores)
            if text_recall_scores else 0.0
        )

        # ── Semantic Similarity: embedding-based relevance ───────────
        # For each turn, computing cosine similarity between the ground_truth_text and each retrieved chunk, then take the max.
        semantic_sim_scores = []
        if retrieved_texts_list and ground_truth_texts_list:
            try:
                from sentence_transformers import SentenceTransformer, util
                _st_model = SentenceTransformer('all-MiniLM-L6-v2')

                for turn_texts, gt_phrase in zip(retrieved_texts_list, ground_truth_texts_list):
                    if not gt_phrase or not turn_texts:
                        continue
                    gt_emb = _st_model.encode(gt_phrase, convert_to_tensor=True)
                    chunk_embs = _st_model.encode(turn_texts, convert_to_tensor=True)
                    sims = util.cos_sim(gt_emb, chunk_embs)[0]
                    semantic_sim_scores.append(float(sims.max()))
            except Exception as e:
                logger.warning(f"Semantic similarity computation failed: {e}")

        avg_semantic_similarity = (
            sum(semantic_sim_scores) / len(semantic_sim_scores)
            if semantic_sim_scores else 0.0
        )

        evidence_results = []
        if retrieved_texts_list and evidence_texts_list:
            for retrieved_texts, evidence_texts in zip(retrieved_texts_list, evidence_texts_list):
                if evidence_texts:
                    evidence_results.append(
                        self._evaluate_evidence_for_turn(retrieved_texts, evidence_texts, k=5)
                    )

        evidence_n = len(evidence_results) if evidence_results else 1
        evidence_turns = len(evidence_results)

        return {
            'avg_coverage': sum(r['coverage'] for r in valid_results) / n,
            'avg_recall_at_5': round(_avg_nested('recall_at_k', 'recall@5'), 4),
            'avg_recall_at_10': round(_avg_nested('recall_at_k', 'recall@10'), 4),
            'avg_recall_at_20': round(_avg_nested('recall_at_k', 'recall@20'), 4),
            'avg_precision_at_5': round(_avg_nested('precision_at_k', 'precision@5'), 4),
            'avg_precision_at_10': round(_avg_nested('precision_at_k', 'precision@10'), 4),
            'avg_precision_at_20': round(_avg_nested('precision_at_k', 'precision@20'), 4),
            'avg_f1_at_5': round(_avg_nested('f1_at_k', 'f1@5'), 4),
            'avg_hit_at_5': round(_avg_nested('hit_at_k', 'hit@5'), 4),
            'avg_hit_at_10': round(_avg_nested('hit_at_k', 'hit@10'), 4),
            'avg_mrr': round(_avg_flat('mrr'), 4),
            'avg_ndcg_at_5': round(_avg_nested('ndcg_at_k', 'ndcg@5'), 4),
            'avg_ndcg_at_10': round(_avg_nested('ndcg_at_k', 'ndcg@10'), 4),
            'avg_map_at_5': round(_avg_nested('average_precision_at_k', 'ap@5'), 4),
            'avg_map_at_10': round(_avg_nested('average_precision_at_k', 'ap@10'), 4),
            'avg_text_recall': round(avg_text_recall, 4),
            'text_recall_turns': len(text_recall_scores),
            'text_recall_hits': sum(text_recall_scores),
            'avg_evidence_recall_at_5': round(sum(r['evidence_recall_at_5'] for r in evidence_results) / evidence_n, 4),
            'avg_evidence_hit_at_5': round(sum(r['evidence_hit_at_5'] for r in evidence_results) / evidence_n, 4),
            'avg_evidence_mrr': round(sum(r['evidence_mrr'] for r in evidence_results) / evidence_n, 4),
            'avg_evidence_ndcg_at_5': round(sum(r['evidence_ndcg_at_5'] for r in evidence_results) / evidence_n, 4),
            'avg_evidence_map_at_5': round(sum(r['evidence_map_at_5'] for r in evidence_results) / evidence_n, 4),
            'avg_evidence_token_f1_at_5': round(sum(r['evidence_token_f1_at_5'] for r in evidence_results) / evidence_n, 4),
            'evidence_turns': evidence_turns,
            'avg_semantic_similarity': round(avg_semantic_similarity, 4),
            'semantic_similarity_turns': len(semantic_sim_scores),
            'total_turns': len(valid_results),
        }

    # ── Comparison ─────────────────────────────────────────────────────

    def _compare_mode_results(self, summaries: Dict[str, Dict]) -> Dict:
        """Compare every ablation mode against the semantic baseline."""
        mode_statuses = {
            mode: {
                "status": summaries.get(mode, {}).get("status", "unknown"),
                "attempted_turns": summaries.get(mode, {}).get("attempted_turns", summaries.get(mode, {}).get("total_turns", 0)),
                "successful_turns": summaries.get(mode, {}).get("total_turns", 0),
                "failed_turns": summaries.get(mode, {}).get("failed_turns", 0),
                "first_error": summaries.get(mode, {}).get("first_error", summaries.get(mode, {}).get("error", "")),
            }
            for mode in BENCHMARK_MODE_ORDER
        }
        baseline = summaries["baseline"]
        full = summaries["sgas"]
        comparison_ready = (
            baseline.get("status") == "completed"
            and full.get("status") == "completed"
            and baseline.get("total_turns", 0) > 0
            and full.get("total_turns", 0) > 0
        )
        if not comparison_ready:
            return {
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "status": "incomplete",
                "reason": "Comparison skipped because baseline or full S-GAS did not complete successfully.",
                "mode_statuses": mode_statuses,
                "baseline_mode": "baseline",
                "mode_order": BENCHMARK_MODE_ORDER,
                "mode_labels": BENCHMARK_MODE_LABELS,
                "modes": {
                    mode: {
                        "label": BENCHMARK_MODE_LABELS[mode],
                        "summary": summaries.get(mode, {}),
                    }
                    for mode in BENCHMARK_MODE_ORDER
                },
                "verdict": {
                    "overall": "Incomplete benchmark; no valid S-GAS effectiveness verdict.",
                },
            }
        comparison = self._compare_results(full, baseline)
        comparison["status"] = "completed"
        comparison["mode_statuses"] = mode_statuses
        comparison["baseline_mode"] = "baseline"
        comparison["mode_order"] = BENCHMARK_MODE_ORDER
        comparison["mode_labels"] = BENCHMARK_MODE_LABELS
        comparison["modes"] = {}

        for mode in BENCHMARK_MODE_ORDER:
            summary = summaries[mode]
            ret = summary.get("retrieval_metrics", {})
            gen = summary.get("generation_metrics", {})
            comparison["modes"][mode] = {
                "label": BENCHMARK_MODE_LABELS[mode],
                "quality": {
                    "recall_at_5": ret.get("avg_recall_at_5", summary.get("avg_retrieval_recall_at_5", 0)),
                    "precision_at_5": ret.get("avg_precision_at_5", summary.get("avg_retrieval_precision_at_5", 0)),
                    "f1_at_5": ret.get("avg_f1_at_5", summary.get("avg_retrieval_f1_at_5", 0)),
                    "hit_at_5": ret.get("avg_hit_at_5", summary.get("avg_retrieval_hit_at_5", 0)),
                    "mrr": ret.get("avg_mrr", summary.get("avg_retrieval_mrr", 0)),
                    "ndcg_at_5": ret.get("avg_ndcg_at_5", summary.get("avg_retrieval_ndcg_at_5", 0)),
                    "map_at_5": ret.get("avg_map_at_5", summary.get("avg_retrieval_map_at_5", 0)),
                    "evidence_recall_at_5": ret.get("avg_evidence_recall_at_5", summary.get("avg_evidence_recall_at_5", 0)),
                    "evidence_hit_at_5": ret.get("avg_evidence_hit_at_5", summary.get("avg_evidence_hit_at_5", 0)),
                    "evidence_mrr": ret.get("avg_evidence_mrr", summary.get("avg_evidence_mrr", 0)),
                    "evidence_ndcg_at_5": ret.get("avg_evidence_ndcg_at_5", summary.get("avg_evidence_ndcg_at_5", 0)),
                    "evidence_map_at_5": ret.get("avg_evidence_map_at_5", summary.get("avg_evidence_map_at_5", 0)),
                    "evidence_token_f1_at_5": ret.get("avg_evidence_token_f1_at_5", summary.get("avg_evidence_token_f1_at_5", 0)),
                    "text_recall": ret.get("avg_text_recall", 0),
                    "retrieval_semantic_similarity": ret.get("avg_semantic_similarity", 0),
                    "answer_semantic_similarity": gen.get("avg_bertscore", summary.get("avg_answer_semantic_similarity", 0)),
                    "answer_token_f1": gen.get("avg_token_f1", summary.get("avg_answer_token_f1", 0)),
                    "answer_exact_match": gen.get("avg_exact_match", summary.get("avg_answer_exact_match", 0)),
                    "answer_rougeL": gen.get("avg_rougeL", summary.get("avg_answer_rougeL", 0)),
                    "multi_turn_accuracy": gen.get("multi_turn_accuracy", 0),
                },
                "performance": {
                    "avg_latency_ms": summary.get("avg_latency_ms", 0),
                    "avg_latency_excl_first_ms": summary.get("avg_latency_excl_first_ms", 0),
                    "avg_vram_gb": summary.get("avg_vram_gb", 0),
                    "peak_vram_gb": summary.get("peak_vram_gb", 0),
                    "avg_cache_hit_rate": summary.get("avg_cache_hit_rate", 0),
                    "total_swap_operations": summary.get("total_swap_operations", 0),
                    "avg_vllm_kv_cache_usage_after": summary.get("avg_vllm_kv_cache_usage_after", 0),
                    "peak_vllm_kv_cache_usage_after": summary.get("peak_vllm_kv_cache_usage_after", 0),
                    "avg_vllm_prefix_cache_hit_rate_delta": summary.get("avg_vllm_prefix_cache_hit_rate_delta", 0),
                    "total_vllm_preemptions": summary.get("total_vllm_preemptions", 0),
                    "avg_vllm_tokens_per_second": summary.get("avg_vllm_tokens_per_second", 0),
                    "avg_vllm_total_tokens_per_second": summary.get("avg_vllm_total_tokens_per_second", 0),
                    "avg_gpu_utilization_pct": summary.get("avg_gpu_utilization_pct", 0),
                    "peak_gpu_utilization_pct": summary.get("peak_gpu_utilization_pct", 0),
                    "peak_gpu_memory_used_mb": summary.get("peak_gpu_memory_used_mb", 0),
                    "avg_graph_nodes": summary.get("avg_graph_nodes", 0),
                    "avg_graph_edges": summary.get("avg_graph_edges", 0),
                },
            }

        return comparison

    def _compare_results(self, sgas_summary: Dict, baseline_summary: Dict) -> Dict:
        """Compare S-GAS vs Baseline summaries and compute deltas."""

        def _delta(sgas_val, baseline_val):
            """Positive delta = S-GAS is better."""
            if baseline_val == 0:
                return {'sgas': sgas_val, 'baseline': baseline_val, 'delta': sgas_val, 'improvement_pct': 0}
            delta = sgas_val - baseline_val
            improvement = (delta / abs(baseline_val)) * 100
            return {
                'sgas': round(sgas_val, 4),
                'baseline': round(baseline_val, 4),
                'delta': round(delta, 4),
                'improvement_pct': round(improvement, 2),
            }

        def _delta_lower_better(sgas_val, baseline_val):
            """For metrics where lower is better (latency, VRAM)."""
            if baseline_val == 0:
                return {'sgas': sgas_val, 'baseline': baseline_val, 'delta': -sgas_val, 'improvement_pct': 0}
            delta = baseline_val - sgas_val  # positive = S-GAS uses less
            improvement = (delta / abs(baseline_val)) * 100
            return {
                'sgas': round(sgas_val, 4),
                'baseline': round(baseline_val, 4),
                'delta': round(delta, 4),
                'improvement_pct': round(improvement, 2),
            }

        # Retrieval quality
        sgas_ret = sgas_summary.get('retrieval_metrics', {})
        base_ret = baseline_summary.get('retrieval_metrics', {})

        # Generation quality
        sgas_gen = sgas_summary.get('generation_metrics', {})
        base_gen = baseline_summary.get('generation_metrics', {})

        comparison = {
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'total_turns': sgas_summary.get('total_turns', 0),
            'quality_metrics': {
                'recall_at_5': _delta(sgas_ret.get('avg_recall_at_5', 0), base_ret.get('avg_recall_at_5', 0)),
                'precision_at_5': _delta(sgas_ret.get('avg_precision_at_5', 0), base_ret.get('avg_precision_at_5', 0)),
                'f1_at_5': _delta(sgas_ret.get('avg_f1_at_5', 0), base_ret.get('avg_f1_at_5', 0)),
                'hit_at_5': _delta(sgas_ret.get('avg_hit_at_5', 0), base_ret.get('avg_hit_at_5', 0)),
                'mrr': _delta(sgas_ret.get('avg_mrr', 0), base_ret.get('avg_mrr', 0)),
                'ndcg_at_5': _delta(sgas_ret.get('avg_ndcg_at_5', 0), base_ret.get('avg_ndcg_at_5', 0)),
                'map_at_5': _delta(sgas_ret.get('avg_map_at_5', 0), base_ret.get('avg_map_at_5', 0)),
                'evidence_recall_at_5': _delta(sgas_ret.get('avg_evidence_recall_at_5', 0), base_ret.get('avg_evidence_recall_at_5', 0)),
                'evidence_hit_at_5': _delta(sgas_ret.get('avg_evidence_hit_at_5', 0), base_ret.get('avg_evidence_hit_at_5', 0)),
                'evidence_mrr': _delta(sgas_ret.get('avg_evidence_mrr', 0), base_ret.get('avg_evidence_mrr', 0)),
                'evidence_ndcg_at_5': _delta(sgas_ret.get('avg_evidence_ndcg_at_5', 0), base_ret.get('avg_evidence_ndcg_at_5', 0)),
                'evidence_map_at_5': _delta(sgas_ret.get('avg_evidence_map_at_5', 0), base_ret.get('avg_evidence_map_at_5', 0)),
                'evidence_token_f1_at_5': _delta(sgas_ret.get('avg_evidence_token_f1_at_5', 0), base_ret.get('avg_evidence_token_f1_at_5', 0)),
                'text_recall': _delta(sgas_ret.get('avg_text_recall', 0), base_ret.get('avg_text_recall', 0)),
                'semantic_similarity': _delta(sgas_ret.get('avg_semantic_similarity', 0), base_ret.get('avg_semantic_similarity', 0)),
                'answer_token_f1': _delta(sgas_gen.get('avg_token_f1', 0), base_gen.get('avg_token_f1', 0)),
                'answer_rougeL': _delta(sgas_gen.get('avg_rougeL', 0), base_gen.get('avg_rougeL', 0)),
                'coverage': _delta(sgas_summary.get('final_coverage', 0), baseline_summary.get('final_coverage', 0)),
                'multi_turn_accuracy': _delta(
                    sgas_gen.get('multi_turn_accuracy', 0),
                    base_gen.get('multi_turn_accuracy', 0),
                ),
            },
            'performance_metrics': {
                'avg_latency_ms': _delta_lower_better(
                    sgas_summary.get('avg_latency_ms', 0),
                    baseline_summary.get('avg_latency_ms', 0),
                ),
                'avg_latency_excl_first_ms': _delta_lower_better(
                    sgas_summary.get('avg_latency_excl_first_ms', 0),
                    baseline_summary.get('avg_latency_excl_first_ms', 0),
                ),
                'avg_vram_gb': _delta_lower_better(
                    sgas_summary.get('avg_vram_gb', 0),
                    baseline_summary.get('avg_vram_gb', 0),
                ),
                'avg_cache_hit_rate': _delta(
                    sgas_summary.get('avg_cache_hit_rate', 0),
                    baseline_summary.get('avg_cache_hit_rate', 0),
                ),
                'avg_vllm_kv_cache_usage_after': _delta_lower_better(
                    sgas_summary.get('avg_vllm_kv_cache_usage_after', 0),
                    baseline_summary.get('avg_vllm_kv_cache_usage_after', 0),
                ),
                'avg_vllm_tokens_per_second': _delta(
                    sgas_summary.get('avg_vllm_tokens_per_second', 0),
                    baseline_summary.get('avg_vllm_tokens_per_second', 0),
                ),
                'avg_gpu_utilization_pct': _delta(
                    sgas_summary.get('avg_gpu_utilization_pct', 0),
                    baseline_summary.get('avg_gpu_utilization_pct', 0),
                ),
                'total_swap_operations': {
                    'sgas': sgas_summary.get('total_swap_operations', 0),
                    'baseline': baseline_summary.get('total_swap_operations', 0),
                },
                'total_vllm_preemptions': {
                    'sgas': sgas_summary.get('total_vllm_preemptions', 0),
                    'baseline': baseline_summary.get('total_vllm_preemptions', 0),
                },
            },
            'graph_usage': {
                'sgas_has_graph': True,
                'baseline_has_graph': False,
                'avg_graph_nodes_sgas': sgas_summary.get('avg_graph_nodes', 0),
                'avg_graph_edges_sgas': sgas_summary.get('avg_graph_edges', 0),
            },
            'verdict': self._generate_verdict(sgas_summary, baseline_summary, sgas_ret, base_ret, sgas_gen, base_gen),
        }

        return comparison

    def _generate_verdict(self, sgas, baseline, sgas_ret, base_ret, sgas_gen, base_gen) -> Dict:
        """Generate a human-readable verdict of which approach wins per metric."""
        verdicts = {}

        # Text-based recall
        s_recall = sgas_ret.get('avg_text_recall') or sgas_ret.get('avg_recall', 0)
        b_recall = base_ret.get('avg_text_recall') or base_ret.get('avg_recall', 0)
        if s_recall > b_recall:
            verdicts['recall'] = f"S-GAS wins ({s_recall:.3f} vs {b_recall:.3f})"
        elif b_recall > s_recall:
            verdicts['recall'] = f"Baseline wins ({b_recall:.3f} vs {s_recall:.3f})"
        else:
            verdicts['recall'] = f"Tie ({s_recall:.3f})"

        # Coverage
        s_cov = sgas.get('final_coverage', 0)
        b_cov = baseline.get('final_coverage', 0)
        if s_cov > b_cov:
            verdicts['coverage'] = f"S-GAS wins ({s_cov:.3f} vs {b_cov:.3f})"
        elif b_cov > s_cov:
            verdicts['coverage'] = f"Baseline wins ({b_cov:.3f} vs {s_cov:.3f})"
        else:
            verdicts['coverage'] = f"Tie ({s_cov:.3f})"

        # Accuracy
        s_acc = sgas_gen.get('multi_turn_accuracy', 0)
        b_acc = base_gen.get('multi_turn_accuracy', 0)
        if s_acc > b_acc:
            verdicts['accuracy'] = f"S-GAS wins ({s_acc * 100:.1f}% vs {b_acc * 100:.1f}%)"
        elif b_acc > s_acc:
            verdicts['accuracy'] = f"Baseline wins ({b_acc * 100:.1f}% vs {s_acc * 100:.1f}%)"
        else:
            verdicts['accuracy'] = f"Tie ({s_acc * 100:.1f}%)"

        # Latency (lower is better; using steady-state avg to exclude warmup turn)
        s_lat = sgas.get('avg_latency_excl_first_ms') or sgas.get('avg_latency_ms', 0)
        b_lat = baseline.get('avg_latency_excl_first_ms') or baseline.get('avg_latency_ms', 0)
        if s_lat < b_lat:
            verdicts['latency'] = f"S-GAS wins ({s_lat:.0f}ms vs {b_lat:.0f}ms)"
        elif b_lat < s_lat:
            verdicts['latency'] = f"Baseline wins ({b_lat:.0f}ms vs {s_lat:.0f}ms)"
        else:
            verdicts['latency'] = f"Tie ({s_lat:.0f}ms)"

        # Overall of GPU
        sgas_wins = sum(1 for v in verdicts.values() if 'S-GAS wins' in v)
        baseline_wins = sum(1 for v in verdicts.values() if 'Baseline wins' in v)
        if sgas_wins > baseline_wins:
            verdicts['overall'] = f"S-GAS algorithm is better ({sgas_wins}/{len(verdicts) - 1} metrics)"
        elif baseline_wins > sgas_wins:
            verdicts['overall'] = f"Baseline is better ({baseline_wins}/{len(verdicts) - 1} metrics)"
        else:
            verdicts['overall'] = "Tie — no clear winner"

        return verdicts

    def _log_comparison(self, comparison: Dict):
        """Logging comparison results to console."""
        logger.info("=" * 60)
        logger.info("BENCHMARK COMPARISON: S-GAS vs Baseline")
        logger.info("=" * 60)

        for cat_name, cat_data in comparison.get('quality_metrics', {}).items():
            if isinstance(cat_data, dict) and 'improvement_pct' in cat_data:
                sign = "+" if cat_data['improvement_pct'] >= 0 else ""
                logger.info(f"  {cat_name}: S-GAS={cat_data['sgas']}, Baseline={cat_data['baseline']} "
                            f"({sign}{cat_data['improvement_pct']}%)")

        for cat_name, cat_data in comparison.get('performance_metrics', {}).items():
            if isinstance(cat_data, dict) and 'improvement_pct' in cat_data:
                sign = "+" if cat_data['improvement_pct'] >= 0 else ""
                logger.info(f"  {cat_name}: S-GAS={cat_data['sgas']}, Baseline={cat_data['baseline']} "
                            f"({sign}{cat_data['improvement_pct']}%)")

        verdict = comparison.get('verdict', {})
        logger.info("-" * 60)
        for k, v in verdict.items():
            logger.info(f"  {k}: {v}")
        logger.info("=" * 60)

    def _collect_benchmark_environment(self) -> Dict[str, Any]:
        """Collect current model/config/hardware info for thesis-ready reports."""
        config_info: Dict[str, Any] = {}
        try:
            from config import Settings
            settings = Settings()
            config_info = {
                "config_path": str(settings.config_path),
                "model_name": settings.get("vllm", {}).get("model_name", ""),
                "max_model_len": settings.get("vllm", {}).get("max_model_len", ""),
                "gpu_memory_utilization": settings.get("vllm", {}).get("gpu_memory_utilization", ""),
                "max_num_seqs": settings.get("vllm", {}).get("max_num_seqs", ""),
                "max_num_batched_tokens": settings.get("vllm", {}).get("max_num_batched_tokens", ""),
                "chunk_size": settings.get("chunking", {}).get("max_chunk_size", ""),
                "chunk_overlap": settings.get("chunking", {}).get("overlap_size", ""),
            }
        except Exception as e:
            config_info = {"error": str(e)}

        gpu_info = ""
        try:
            result = subprocess.run(
                ["nvidia-smi", "--query-gpu=name,memory.total,driver_version", "--format=csv,noheader"],
                capture_output=True,
                text=True,
                timeout=3,
                check=False,
            )
            gpu_info = result.stdout.strip() if result.returncode == 0 else ""
        except Exception:
            gpu_info = ""

        return {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "platform": platform.platform(),
            "processor": platform.processor(),
            "python_version": platform.python_version(),
            "gpu": gpu_info or "NVIDIA GPU info unavailable",
            "config": config_info,
        }

    @staticmethod
    def _docx_add_kv_table(document, title: str, rows: List[tuple]) -> None:
        document.add_heading(title, level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Parameter"
        hdr[1].text = "Value"
        for key, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(value)

    def _generate_docx_summary_report(
        self,
        scenario,
        mode_results: Dict[str, Dict],
        comparison: Dict,
        output_path: Path,
    ) -> None:
        """Generate a compact DOCX report with average results for all modes."""
        try:
            from docx import Document
        except ImportError as e:
            logger.warning(f"DOCX report skipped: python-docx is unavailable: {e}")
            return

        env = self._collect_benchmark_environment()
        config = env.get("config", {})

        document = Document()
        document.add_heading("S-GAS Benchmark Summary Report", level=1)
        document.add_paragraph(f"Scenario: {scenario.name}")
        document.add_paragraph(f"Description: {scenario.description}")
        document.add_paragraph(f"Document: {scenario.document}")
        document.add_paragraph(f"Generated at: {env.get('generated_at')}")

        self._docx_add_kv_table(document, "Current Setup", [
            ("Config path", config.get("config_path", "")),
            ("Model", config.get("model_name", "")),
            ("Max model length", config.get("max_model_len", "")),
            ("GPU memory utilization", config.get("gpu_memory_utilization", "")),
            ("Max sequences", config.get("max_num_seqs", "")),
            ("Max batched tokens", config.get("max_num_batched_tokens", "")),
            ("Chunk size", config.get("chunk_size", "")),
            ("Chunk overlap", config.get("chunk_overlap", "")),
            ("Platform", env.get("platform", "")),
            ("Processor", env.get("processor", "")),
            ("Python", env.get("python_version", "")),
            ("GPU", env.get("gpu", "")),
        ])

        document.add_heading("Average Results by Algorithm Mode", level=2)
        table = document.add_table(rows=1, cols=12)
        table.style = "Table Grid"
        headers = [
            "Mode", "Evidence Recall@5", "Recall@5", "Answer Token-F1",
            "ROUGE-L", "Avg Latency ms", "Avg VRAM GB", "KV Cache Avg",
            "GPU Util Avg %", "Generated tok/s", "Swaps", "Preemptions",
        ]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        for mode in BENCHMARK_MODE_ORDER:
            result = mode_results.get(mode, {})
            summary = result.get("summary", {})
            retrieval = summary.get("retrieval_metrics", {})
            generation = summary.get("generation_metrics", {})
            values = [
                result.get("label", BENCHMARK_MODE_LABELS.get(mode, mode)),
                retrieval.get("avg_evidence_recall_at_5", summary.get("avg_evidence_recall_at_5", 0)),
                retrieval.get("avg_recall_at_5", summary.get("avg_retrieval_recall_at_5", 0)),
                generation.get("avg_token_f1", summary.get("avg_answer_token_f1", 0)),
                generation.get("avg_rougeL", summary.get("avg_answer_rougeL", 0)),
                summary.get("avg_latency_ms", 0),
                summary.get("avg_vram_gb", 0),
                summary.get("avg_vllm_kv_cache_usage_after", 0),
                summary.get("avg_gpu_utilization_pct", 0),
                summary.get("avg_vllm_tokens_per_second", 0),
                summary.get("total_swap_operations", 0),
                summary.get("total_vllm_preemptions", 0),
            ]
            cells = table.add_row().cells
            for idx, value in enumerate(values):
                if isinstance(value, float):
                    value = round(value, 4)
                cells[idx].text = str(value)

        verdict = comparison.get("verdict", {})
        self._docx_add_kv_table(document, "Full S-GAS vs Baseline Verdict", [
            ("Overall", verdict.get("overall", "")),
            ("Recall", verdict.get("recall", "")),
            ("Coverage", verdict.get("coverage", "")),
            ("Accuracy", verdict.get("accuracy", "")),
            ("Latency", verdict.get("latency", "")),
        ])

        document.add_paragraph(
            "Note: Evidence metrics are chunk-ID-independent and are recommended as "
            "primary retrieval-quality indicators when chunk size differs between configs."
        )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))

    # ── Session helpers ────────────────────────────────────────────────

    def _create_session(self, sessions: dict) -> str:
        session_id = f"session_{uuid.uuid4().hex[:12]}"
        sessions[session_id] = {
            'iteration': 0,
            'excluded_chunks': [],
            'created_at': datetime.now(timezone.utc).isoformat(),
            'swap_initialized': False,
            'graph_initialized': False,
        }
        return session_id

    def _cleanup_session(self, session_id: str) -> None:
        """Removing a finished benchmark session's vector store chunks and in-memory state."""
        try:
            if self._chat_service is not None:
                self._chat_service._vector_store.delete_session_chunks(session_id)
                logger.info(f"Cleaned up vector store for session {session_id}")
        except Exception as e:
            logger.warning(f"Could not clean up session {session_id}: {e}")

    def _save_graph_snapshot(self, scenario_name: str, timestamp: str) -> None:
        """Persisting the current S-GAS knowledge graph to disk before it is reset."""
        if self._chat_service is None:
            return
        try:
            snapshot = self._chat_service.get_graph_snapshot()
            node_count = len(snapshot.get('nodes', []))
            if node_count == 0:
                logger.warning("Graph snapshot skipped: graph is empty after S-GAS run")
                return

            # Always-current alias (used by the web endpoint as fallback)
            latest_path = self.results_dir / "latest_sgas_graph.json"
            with open(latest_path, 'w', encoding='utf-8') as f:
                json.dump(snapshot, f, ensure_ascii=False, default=str)

            # Timestamped archive copy
            archive_path = self.results_dir / f"{scenario_name}_graph_{timestamp}.json"
            with open(archive_path, 'w', encoding='utf-8') as f:
                json.dump(snapshot, f, ensure_ascii=False, default=str)

            logger.info(
                f"Graph snapshot saved: {node_count} nodes, "
                f"{len(snapshot.get('edges', []))} edges → {latest_path}"
            )
        except Exception as e:
            logger.warning(f"Graph snapshot save failed: {e}")

    def _upload_document(self, session_id: str, doc_path: Path, document_processor,
                         chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None):
        """Uploading a document for benchmarking."""
        chunker = getattr(document_processor, '_chunker', None)
        orig_size = orig_overlap = None
        if chunker is not None and chunk_size is not None:
            orig_size = getattr(chunker, '_max_chunk_size', None)
            orig_overlap = getattr(chunker, '_overlap_size', None)
            chunker._max_chunk_size = chunk_size
            chunker._overlap_size = chunk_overlap if chunk_overlap is not None else max(1, chunk_size // 8)
            logger.info(f"Benchmark chunk size patched: {orig_size}→{chunk_size} words")
        try:
            result = document_processor.process_document(
                doc_path, session_id=session_id, metadata={"document_type": "general"},
            )
        finally:
            if chunker is not None and orig_size is not None:
                chunker._max_chunk_size = orig_size
                chunker._overlap_size = orig_overlap
        if result.status != 'success':
            raise RuntimeError(f"Document upload failed: {result.error}")
