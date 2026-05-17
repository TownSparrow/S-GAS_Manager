from pathlib import Path
from typing import Dict, Any
import logging

from app.interfaces.document_loader import IDocumentLoader

logger = logging.getLogger(__name__)


class DOCXLoader(IDocumentLoader):
    def load(self, file_path: Path) -> Dict[str, Any]:
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            text = "\n".join(text_parts)
            if not text.strip():
                raise ValueError("DOCX has no extractable text")
            return {'text': text}
        except ImportError:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
