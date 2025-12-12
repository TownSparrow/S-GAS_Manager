from pathlib import Path
from typing import List, Dict, Any
import logging
from abc import ABC, abstractmethod

logger = logging.getLogger(__name__)

class DocumentLoader(ABC):
    """Base class for document loaders"""
    
    @abstractmethod
    def load(self, file_path: Path) -> Dict[str, Any]:
        pass


class PDFLoader(DocumentLoader):
    def load(self, file_path: Path) -> Dict[str, Any]:
        try:
            import PyPDF2
            
            text = ""
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                logger.info(f"📄 Loading PDF with {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract page {page_num}: {e}")
            
            if not text.strip():
                logger.warning(f"⚠️ No text extracted from PDF")
                raise ValueError("PDF has no extractable text")
            
            logger.info(f"✅ Extracted {len(text)} characters from PDF")
            
            return {'text': text}
        
        except ImportError:
            raise RuntimeError("❌ PyPDF2 not installed. Install with: pip install PyPDF2")
        except Exception as e:
            logger.error(f"❌ Failed to load PDF: {e}")
            raise


class TextLoader(DocumentLoader):
    def load(self, file_path: Path) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if not text.strip():
                logger.warning(f"⚠️ Text file is empty")
            
            logger.info(f"✅ Loaded {len(text)} characters from text file")
            
            return {'text': text}
        
        except UnicodeDecodeError:
            logger.warning(f"⚠️ UTF-8 decoding failed, trying latin-1")
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return {'text': text}
            except Exception as e:
                logger.error(f"❌ Failed to load text file: {e}")
                raise
        except Exception as e:
            logger.error(f"❌ Failed to load text file: {e}")
            raise


class DOCXLoader(DocumentLoader):
    def load(self, file_path: Path) -> Dict[str, Any]:
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                logger.warning(f"⚠️ No text extracted from DOCX")
                raise ValueError("DOCX has no extractable text")
            
            logger.info(f"✅ Extracted {len(text)} characters from DOCX")
            
            return {'text': text}
        
        except ImportError:
            raise RuntimeError("❌ python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"❌ Failed to load DOCX: {e}")
            raise